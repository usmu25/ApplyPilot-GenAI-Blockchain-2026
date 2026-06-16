import os
import re
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt


def safe_filename(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")[:50]


def convert_docx_to_pdf(docx_path: str) -> str:
    output_dir = os.path.dirname(docx_path)

    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE
    )

    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF conversion failed: {pdf_path}")

    return pdf_path


def add_clean_text_to_doc(doc: Document, text: str):
    for line in text.splitlines():
        clean_line = line.strip()

        if not clean_line:
            doc.add_paragraph("")
            continue

        normalized = clean_line.replace("**", "").strip()

        is_heading = (
            normalized.endswith(":")
            or normalized.lower() in [
                "profile",
                "research profile",
                "contact information",
                "education",
                "research experience",
                "professional experience",
                "technical skills",
                "selected projects",
                "projects",
                "publications",
                "academic strengths",
                "languages",
                "certifications"
            ]
        )

        if is_heading:
            doc.add_heading(normalized.replace(":", ""), level=2)
        elif normalized.startswith("-") or normalized.startswith("•") or normalized.startswith("*"):
            doc.add_paragraph(normalized.lstrip("-•* ").strip(), style="List Bullet")
        else:
            doc.add_paragraph(normalized)


def save_cv_as_docx(cv_text: str, opportunity: dict, opportunity_number: str) -> str:
    os.makedirs("data/final_cvs", exist_ok=True)

    title = opportunity.get("title", "opportunity")

    filename = (
        f"CV_opportunity_{opportunity_number}_"
        f"{safe_filename(title)}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    file_path = os.path.join("data/final_cvs", filename)

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    add_clean_text_to_doc(doc, cv_text)

    doc.save(file_path)
    return file_path


def save_cv_as_pdf(cv_text: str, opportunity: dict, opportunity_number: str) -> str:
    docx_path = save_cv_as_docx(cv_text, opportunity, opportunity_number)
    return convert_docx_to_pdf(docx_path)


def save_cover_letter_as_docx(cover_letter_text: str, opportunity: dict, opportunity_number: str) -> str:
    os.makedirs("data/cover_letters", exist_ok=True)

    title = opportunity.get("title", "opportunity")

    filename = (
        f"Cover_Letter_opportunity_{opportunity_number}_"
        f"{safe_filename(title)}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    file_path = os.path.join("data/cover_letters", filename)

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    add_clean_text_to_doc(doc, cover_letter_text)

    doc.save(file_path)
    return file_path


def save_cover_letter_as_pdf(cover_letter_text: str, opportunity: dict, opportunity_number: str) -> str:
    docx_path = save_cover_letter_as_docx(
        cover_letter_text=cover_letter_text,
        opportunity=opportunity,
        opportunity_number=opportunity_number
    )

    return convert_docx_to_pdf(docx_path)
